"""Generate project documentation as .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1a, 0x52, 0x76)
MID_BLUE  = RGBColor(0x2e, 0x86, 0xc1)
GRAY      = RGBColor(0x44, 0x44, 0x44)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(20)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(15)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

# ── TITLE PAGE ───────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("IEU-CEPA Compliance Scanner")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Projektdokumentation – Prototyp")
run2.font.size = Pt(14)
run2.font.color.rgb = GRAY

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Stand: März 2026  |  Version 1.0  |  HRW Master Projekt 2")

doc.add_page_break()

# ── 1. ÜBERBLICK ─────────────────────────────────────────────────────────────
h1("1. Projektüberblick")
body(
    "IEU-CEPA Compliance Scanner ist eine mobile App (React Native / Expo) mit einem "
    "Python-Backend, die indonesischen KMU-Exporteuren hilft, EU-Compliance-Anforderungen "
    "im Rahmen des IEU-CEPA Abkommens (Indonesia–EU Comprehensive Economic Partnership Agreement) "
    "zu verstehen und zu erfüllen."
)
body(
    "Das Backend nutzt ein RAG-System (Retrieval-Augmented Generation): Offizielle EU-Rechtsdokumente "
    "werden als Wissensbasis indexiert. Nutzerfragen und Compliance-Analysen werden mit diesem "
    "Wissen angereichert, bevor sie an das KI-Modell (GPT-4o-mini) gesendet werden."
)

h2("Projektstruktur")
code("ieu-cepa-scan/")
code("├── mobile/          ← React Native Expo App")
code("│   ├── app/         ← Screens (Expo Router)")
code("│   ├── components/  ← Wiederverwendbare Komponenten")
code("│   ├── context/     ← AuthContext")
code("│   ├── lib/         ← Storage & API Utilities")
code("│   └── constants/   ← API-Konfiguration")
code("├── backend/         ← Python FastAPI + RAG")
code("│   ├── rag/")
code("│   │   ├── ingest.py    ← PDFs indexieren")
code("│   │   └── retriever.py ← RAG-Abruf + OpenAI")
code("│   ├── main.py          ← API Endpoints")
code("│   ├── documents/       ← EU-Rechtsdokumente (PDFs)")
code("│   ├── chroma_db/       ← Vektordatenbank (auto-generiert)")
code("│   └── venv/            ← Python virtual environment")
code("└── docs/                ← Projektdokumentation")

doc.add_page_break()

# ── 2. MOBILE APP ─────────────────────────────────────────────────────────────
h1("2. Mobile App")

h2("Technologie-Stack")
add_table(
    ["Bereich", "Technologie"],
    [
        ["Mobile Framework", "React Native + Expo 54"],
        ["Routing", "Expo Router (File-based)"],
        ["Icons", "@expo/vector-icons (Ionicons)"],
        ["Lokale Datenspeicherung", "AsyncStorage"],
        ["Sprache", "TypeScript"],
        ["Styling", "React Native StyleSheet"],
    ]
)

h2("Funktionsübersicht")
add_table(
    ["Funktion", "Status", "Beschreibung"],
    [
        ["Registrierung / Login", "✅ Fertig", "Lokal via AsyncStorage, kein Server"],
        ["Dashboard", "✅ Fertig", "Letzter Scan + Feature-Karten + Chat-Zugang"],
        ["Readiness Scan", "✅ Fertig", "4-Schritt-Formular + KI-Analyse via Backend"],
        ["Aktionsplan", "✅ Fertig", "Priorisierte Aufgabenliste mit Fortschrittsanzeige"],
        ["KI-Assistent", "✅ Fertig", "Chat mit RAG-Wissensbasis (EU-Rechtsdokumente)"],
        ["Profil & Verlauf", "✅ Fertig", "Nutzerinfo, Statistiken, letzte 3 Scans"],
        ["Learning Hub", "🔜 Geplant", "Platzhalter mit 6 geplanten Modulen"],
        ["Community Hub", "🔜 Geplant", "Peer-to-Peer Austausch"],
        ["Expert Matching", "🔜 Geplant", "Verbindung mit Compliance-Experten"],
    ]
)

h2("Screens im Detail")

h3("Authentifizierung (app/auth.tsx)")
bullet("Login-Tab: E-Mail + Passwort")
bullet("Registrierungs-Tab: Name, E-Mail, Passwort, Unternehmen, Branche")
bullet("Validierung: E-Mail-Format, Passwort min. 6 Zeichen, Pflichtfelder")
bullet("Verfügbare Branchen: Herstellung, Textil, Lebensmittel, Elektronik, Möbel, Chemikalien, Landwirtschaft, Sonstiges")
bullet("Wichtig: Kein Server – alle Daten lokal im AsyncStorage des Geräts")

h3("Dashboard (app/(tabs)/dashboard.tsx)")
bullet("Begrüßung mit Nutzername und Unternehmen")
bullet("Letzter Scan: Score, Risikolevel, Produktname, Link zum Aktionsplan")
bullet("Leerer Zustand: Aufforderung zum ersten Scan falls noch keiner vorhanden")
bullet("Feature-Karten (2-spaltig): Readiness Scan, Aktionsplan, KI-Assistent, Learning Hub")
bullet("KI-Assistent-Button öffnet das Chat-Modal")

h3("Readiness Scan (app/(tabs)/scan.tsx)")
body("Vierschrittiger Assistent zur Compliance-Bewertung:")
bullet("Schritt 1 – Unternehmensprofil: Name, Branche, Größe (Mikro/Klein/Mittel/Groß), Region")
bullet("Schritt 2 – Produkt & Markt: Produkttyp, HS-Code, Zielland in der EU, Exporterfahrung")
bullet("Schritt 3 – Compliance-Selbsteinschätzung:")
bullet("Digital Product Passport (DPP) – EU 2024/1781", level=1)
bullet("EUDR (Entwaldungsverordnung) – EU 2023/1115", level=1)
bullet("CE-Kennzeichnung – EU 768/2008", level=1)
bullet("ESG-Berichterstattung (CSRD) – EU 2022/2464", level=1)
bullet("Ursprungsnachweis (Präferenzzölle)", level=1)
bullet("Lebensmittelsicherheit – EU 178/2002", level=1)
bullet("Schritt 4 – Ergebnisse: Score-Kreis, Risikolevel, erfüllte/fehlende Anforderungen, Aktionsplan")
body("Score-Farben: grün (≥70), orange (40–69), rot (<40)")
body("Ergebnisse werden automatisch lokal gespeichert.")

h3("Aktionsplan (app/(tabs)/action-plan.tsx)")
bullet("Aufgaben aus dem letzten Scan werden angezeigt")
bullet("Filterbar nach Priorität: Alle / Hoch / Mittel / Niedrig")
bullet("Checkboxen zum Abhaken erledigter Aufgaben")
bullet("Fortschrittsbalken: Erledigt / Gesamt in Prozent")
bullet("Prioritätszuweisung: obere 34% = Hoch, 34–67% = Mittel, Rest = Niedrig")

h3("KI-Assistent (components/ChatModal.tsx)")
bullet("Schnellaktionen: 'Was ist der DPP?', 'EUDR erklären', 'CE-Kennzeichnung', 'ESG-Pflichten'")
bullet("Nachrichten: Nutzer rechts (dunkelblau), KI links (grau)")
bullet("Quellendokumente als Badge unter jeder KI-Antwort")
bullet("Gesprächsverlauf wird innerhalb der Sitzung gehalten")
bullet("API-Call: POST /chat an das lokale Backend")

h3("Profil (app/(tabs)/profile.tsx)")
bullet("Nutzerdaten: Name, E-Mail, Unternehmen, Branche")
bullet("Statistiken: Anzahl Scans, Aktionspunkte, letzter Score")
bullet("Scan-Verlauf: letzte 3 Scans mit Datum, Produkt, Score")
bullet("Einstellungen: Sprache (Deutsch), Benachrichtigungen")
bullet("App-Info: Version, Backend-URL, Datenspeicherung")
bullet("Logout mit Bestätigungsdialog")

h2("Datenspeicherung")
body("Alle Daten liegen lokal auf dem Gerät (kein Cloud-Speicher):")
add_table(
    ["Schlüssel", "Inhalt"],
    [
        ["@ieu_cepa:user", "Nutzerprofil: Name, E-Mail, Unternehmen, Branche, ID"],
        ["@ieu_cepa:scans", "Liste aller Scan-Ergebnisse mit Score, Risikolevel, Anforderungen, Aktionsplan"],
    ]
)

h2("API-Konfiguration")
body("Konfigurationsdatei: mobile/constants/api.ts")
code('// Für Emulator:    "http://localhost:8000"')
code('// Für echtes Handy: IP des PCs im gleichen WLAN')
code('export const API_BASE_URL = "http://192.168.178.180:8000";')
body("Handy und PC müssen im gleichen WLAN-Netzwerk sein.")

doc.add_page_break()

# ── 3. BACKEND ────────────────────────────────────────────────────────────────
h1("3. Backend")

h2("Technologie-Stack")
add_table(
    ["Bereich", "Technologie"],
    [
        ["Framework", "Python FastAPI"],
        ["RAG-Framework", "LangChain"],
        ["Vektordatenbank", "ChromaDB (lokal, persistent)"],
        ["Embeddings", "OpenAI text-embedding-3-small"],
        ["KI-Modell", "OpenAI GPT-4o-mini"],
        ["API-Key", "Windows Systemvariable OPENAI_API_KEY"],
        ["Laufzeitumgebung", "Python venv (backend/venv/)"],
    ]
)

h2("Backend starten")
code("cd backend")
code("venv/Scripts/uvicorn main:app --reload --host 0.0.0.0")
body("Erreichbar unter: http://192.168.178.180:8000")
body("Docs (Swagger UI): http://192.168.178.180:8000/docs")

h2("API Endpoints")

h3("GET /health")
body("Prüft ob das Backend läuft.")
code('Response: { "status": "ok" }')

h3("POST /chat")
body("KI-Antwort auf eine Compliance-Frage, basierend auf den indizierten EU-Rechtsdokumenten.")
body("Request:")
code('{ "message": "Was ist der Digital Product Passport?",')
code('  "history": [{ "role": "user", "content": "..." }, ...] }')
body("Response:")
code('{ "answer": "Der DPP ist eine digitale Dokumentation...",')
code('  "sources": ["DPP/OJ_L_202401781.pdf", "DPP/zibold.pdf"] }')

h3("POST /analyze-readiness")
body("Analysiert Compliance-Daten eines Unternehmens und erstellt einen Readiness-Score mit Aktionsplan.")
body("Request-Felder:")
add_table(
    ["Feld", "Typ", "Beschreibung"],
    [
        ["company_name", "string", "Unternehmensname"],
        ["sector", "string", "Branche"],
        ["company_size", "string", "Mikro / Klein / Mittel / Groß"],
        ["product_type", "string", "Produktbeschreibung"],
        ["target_country", "string", "Zielland in der EU"],
        ["compliance_eudr", "boolean", "EUDR-Sorgfaltspflicht erfüllt?"],
        ["compliance_dpp", "boolean", "DPP vorhanden?"],
        ["compliance_ce", "boolean", "CE-Kennzeichnung vorhanden?"],
        ["compliance_esg", "boolean", "ESG-Berichterstattung?"],
        ["compliance_origin", "boolean", "Ursprungsnachweis?"],
        ["compliance_food_safety", "boolean", "Lebensmittelsicherheit?"],
    ]
)
body("Response-Felder:")
add_table(
    ["Feld", "Typ", "Beschreibung"],
    [
        ["readiness_score", "integer", "0–100, KI-gewichteter Score"],
        ["risk_level", "string", "Niedrig / Mittel / Hoch"],
        ["missing_requirements", "list[string]", "Fehlende Anforderungen"],
        ["action_plan", "list[string]", "Priorisierter Aktionsplan (max. 6 Schritte)"],
        ["analysis", "string", "Freitextanalyse durch KI"],
    ]
)

doc.add_page_break()

# ── 4. RAG-SYSTEM ─────────────────────────────────────────────────────────────
h1("4. RAG-System")

h2("Funktionsweise")
body(
    "RAG (Retrieval-Augmented Generation) kombiniert klassische Dokumentensuche mit KI-Textgenerierung. "
    "Anstatt nur auf das trainierte Wissen von GPT-4o-mini zu vertrauen, werden bei jeder Anfrage "
    "zuerst die relevantesten Textstellen aus den offiziellen EU-Rechtsdokumenten abgerufen "
    "und dem Modell als Kontext mitgegeben."
)

h3("Ablauf – Einmalig (Dokumente indexieren)")
bullet("Alle PDFs aus backend/documents/ (inkl. Unterordner) werden geladen")
bullet("Jedes Dokument wird in Chunks aufgeteilt (1000 Zeichen, 200 Zeichen Überlappung)")
bullet("Jeder Chunk wird mit OpenAI text-embedding-3-small in einen Vektor (Zahlenreihe) umgewandelt")
bullet("Vektoren werden in ChromaDB auf der lokalen Festplatte gespeichert")
body("Befehl:")
code("cd backend")
code("venv/Scripts/python -m rag.ingest")

h3("Ablauf – Pro Nutzeranfrage")
bullet("Nutzerfrage wird ebenfalls in einen Vektor umgewandelt")
bullet("Die 5 ähnlichsten Dokument-Chunks werden aus ChromaDB abgerufen")
bullet("Prompt wird zusammengestellt: Systemrolle + abgerufene Chunks als Kontext + Frage")
bullet("GPT-4o-mini generiert eine Antwort basierend auf diesem Kontext")
bullet("Antwort + Quelldokumente werden zurückgegeben")

h2("Systemrolle (System Prompt)")
body(
    "Das KI-Modell wird als 'EU-Compliance-Experte für indonesische KMU-Exporteure' positioniert. "
    "Es wird angewiesen, Antworten ausschließlich auf Basis der bereitgestellten Quelldokumente zu geben, "
    "und in der Sprache des Nutzers zu antworten (Deutsch, Englisch oder Indonesisch)."
)

h2("Konfiguration")
add_table(
    ["Parameter", "Wert"],
    [
        ["Embedding-Modell", "text-embedding-3-small"],
        ["Chat-Modell", "gpt-4o-mini"],
        ["Chunk-Größe", "1000 Zeichen"],
        ["Chunk-Überlappung", "200 Zeichen"],
        ["Abgerufene Chunks (Top-K)", "5"],
        ["Temperatur (Kreativität)", "0.2 (sehr deterministisch)"],
        ["Vektordatenbank", "ChromaDB, lokal unter backend/chroma_db/"],
        ["Collection Name", "eu_compliance_docs"],
    ]
)

h2("Kosten (OpenAI API)")
add_table(
    ["Vorgang", "Modell", "Kosten"],
    [
        ["Ingest (einmalig, alle PDFs)", "text-embedding-3-small", "~$0.07"],
        ["Chat-Nachricht (Nutzer)", "gpt-4o-mini", "~$0.001 pro Nachricht"],
        ["Readiness-Scan", "gpt-4o-mini", "~$0.001 pro Scan"],
    ]
)

h2("Indexierte EU-Dokumente")
add_table(
    ["Kategorie", "Dokumente"],
    [
        ["IEU-CEPA", "Kapitel 1–25 des EU-Indonesia Handelsabkommens"],
        ["Umwelt & Nachhaltigkeit", "EUDR, DPP, ESPR, CSRD, CSDDD, PPWR"],
        ["Chemikalien & Materialien", "REACH, RoHS, EU Battery Regulation"],
        ["Lebensmittel", "EU Food Safety Regulation (178/2002)"],
        ["Digital & Daten", "GDPR, EU Digital Services Act"],
    ]
)

doc.add_page_break()

# ── 5. AUTHENTIFIZIERUNG ──────────────────────────────────────────────────────
h1("5. Authentifizierung & Datenschutz")

h2("Aktuelle Implementierung (Prototyp)")
body(
    "Die Authentifizierung ist vollständig lokal implementiert – es gibt keinen Authentifizierungsserver. "
    "Nutzerdaten (Name, E-Mail, Passwort-Hash) werden ausschließlich im AsyncStorage des Geräts gespeichert."
)
bullet("Kein Passwort-Server: Login prüft nur ob die E-Mail bereits registriert ist")
bullet("Kein Passwort-Hash: Passwort wird nicht dauerhaft gespeichert")
bullet("Daten gehen verloren bei App-Deinstallation oder Cache-Leerung")
bullet("Geeignet für: Prototyp, Demos, persönliche Tests")

h2("Produktionsreife Lösung (Ausblick)")
bullet("Server-seitige Authentifizierung mit JWT-Tokens")
bullet("Passwort-Hashing (bcrypt)")
bullet("Nutzerdaten in PostgreSQL-Datenbank")
bullet("HTTPS/TLS für alle API-Verbindungen")
bullet("Backend-Hosting auf Cloud-Server (z.B. Railway, Render, Fly.io)")

h2("OpenAI API-Key Sicherheit")
body("Der API-Key wird als Windows-Systemvariable gespeichert (nicht in Dateien):")
code("[System.Environment]::SetEnvironmentVariable(\"OPENAI_API_KEY\", \"sk-...\", \"User\")")
bullet("Key liegt in der Windows-Registry, nicht im Projektordner")
bullet("Wird nie in Git eingecheckt")
bullet("Python liest ihn automatisch via os.getenv('OPENAI_API_KEY')")
bullet("Empfehlung: Monatliches Limit auf platform.openai.com setzen (z.B. $5)")

doc.add_page_break()

# ── 6. SETUP ──────────────────────────────────────────────────────────────────
h1("6. Setup & Installation")

h2("Voraussetzungen")
bullet("Node.js (v18+)")
bullet("Python 3.9+")
bullet("Expo Go App auf dem Handy (iOS App Store / Google Play)")
bullet("OpenAI API-Key (platform.openai.com)")

h2("Backend Setup")
code("cd backend")
code("python -m venv venv")
code("venv/Scripts/pip install -r requirements.txt")
body("OpenAI Key als Windows-Systemvariable setzen (PowerShell):")
code("[System.Environment]::SetEnvironmentVariable(\"OPENAI_API_KEY\", \"sk-...\", \"User\")")
body("PDFs in backend/documents/ ablegen, dann indexieren:")
code("venv/Scripts/python -m rag.ingest")
body("Backend starten:")
code("venv/Scripts/uvicorn main:app --reload --host 0.0.0.0")

h2("Mobile App Setup")
code("cd mobile")
code("npm install")
code("npx expo start --clear")
body("QR-Code mit Expo Go App scannen. Handy und PC müssen im gleichen WLAN sein.")
body("Bei echtem Handy: IP des PCs in mobile/constants/api.ts eintragen.")

h2("Bekannte Einschränkungen (Prototyp)")
bullet("localhost:8000 funktioniert nur auf Emulator, nicht auf echtem Handy")
bullet("Keine Offline-Nutzung der KI-Funktionen (Backend muss laufen)")
bullet("Kein Multi-User-System (jedes Gerät hat eigene lokale Daten)")
bullet("Learning Hub noch nicht implementiert")

doc.add_page_break()

# ── 7. ROADMAP ────────────────────────────────────────────────────────────────
h1("7. Roadmap")

add_table(
    ["Feature", "Priorität", "Aufwand", "Beschreibung"],
    [
        ["Learning Hub", "Hoch", "Mittel", "Video-Microlearning Module zu EU-Regularien"],
        ["Echte Authentifizierung", "Hoch", "Mittel", "Server-seitiges Auth mit JWT, PostgreSQL"],
        ["Backend-Hosting", "Hoch", "Gering", "Deployment auf Railway / Render / Fly.io"],
        ["Mehrsprachigkeit", "Mittel", "Mittel", "Englisch + Bahasa Indonesia"],
        ["Push-Benachrichtigungen", "Mittel", "Gering", "Erinnerungen für Aktionspunkte"],
        ["Community Hub", "Niedrig", "Hoch", "Peer-to-Peer Austausch zwischen Exporteuren"],
        ["Expert Matching", "Niedrig", "Hoch", "Verbindung mit Compliance-Experten"],
        ["Dokumenten-Upload", "Mittel", "Mittel", "Analyse eigener Unternehmensdokumente"],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
import os
out = os.path.join(os.path.dirname(__file__), "IEU-CEPA_Compliance_Scanner_Dokumentation.docx")
doc.save(out)
print(f"Saved: {out}")
